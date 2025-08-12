from docx import Document
import re, os
from typing import Dict, List, Any, Union
from datetime import datetime
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self):
        self.supported_formats = ['.docx']
        self.nlp = None
        self._load_nlp_model()

        self.section_patterns = {
            'article': r'article\s+(\d+)\.?\s*(.+)',
            'clause': r'clause\s+(\d+)\.?\s*(.+)',
            'section': r'section\s+(\d+)\.?\s*(.+)',
            'part': r'part\s+([ivxlcdm]+|\d+)\.?\s*(.+)',
            'chapter': r'chapter\s+(\d+)\.?\s*(.+)',
            'whereas': r'whereas\s*[,;:]?\s*(.+)',
            'resolved': r'(?:resolved|it is resolved)\s*[,;:]?\s*(.+)',
            'definitions': r'definitions?\s*[,;:]?\s*(.+)?',
            'schedule': r'schedule\s+([A-Z]|\d+)\s*(.+)',
            'appendix': r'appendix\s+([A-Z]|\d+)\s*(.+)',
            'recital': r'recital\s+([A-Z]|\d+)\s*(.+)'
        }

        self.type_indicators = {
            'articles_of_association': {
                'primary': ['articles of association', 'company governance', 'board of directors', 'shareholders'],
                'secondary': ['quorum', 'general meeting', 'voting rights', 'share capital', 'dividend distribution'],
                'patterns': [r'article\s+\d+', r'board\s+of\s+directors', r'general\s+meeting', r'shareholders?\s+meeting']
            },
            'memorandum_of_association': {
                'primary': ['memorandum of association', 'company objects', 'liability limited', 'registered office'],
                'secondary': ['authorized capital', 'company purposes', 'business activities', 'share classes'],
                'patterns': [r'objects?\s+of\s+(?:the\s+)?company', r'liability.*limited', r'registered\s+office']
            },
            'incorporation_application': {
                'primary': ['incorporation application', 'company registration', 'registration authority', 'ADGM registration'],
                'secondary': ['proposed company name', 'registered address', 'initial directors', 'application form'],
                'patterns': [r'application\s+for\s+incorporation', r'proposed\s+(?:company\s+)?name', r'registration\s+authority']
            },
            'board_resolution': {
                'primary': ['board resolution', 'directors resolution', 'board meeting', 'corporate resolution'],
                'secondary': ['resolved that', 'unanimously resolved', 'board unanimously', 'meeting minutes'],
                'patterns': [r'resolved\s+that', r'board\s+resolution', r'directors?\s+resolve', r'unanimously\s+resolved']
            },
            'register_members': {
                'primary': ['register of members', 'register of directors', 'membership register', 'shareholder register'],
                'secondary': ['member details', 'share certificates', 'shareholding', 'director information'],
                'patterns': [r'register\s+of\s+(?:members|directors)', r'member\s+(?:name|details)', r'shares?\s+held']
            },
            'employment_contract': {
                'primary': ['employment contract', 'employment agreement', 'service agreement', 'employment terms'],
                'secondary': ['salary', 'working hours', 'termination', 'probation', 'job description'],
                'patterns': [r'employment\s+(?:contract|agreement)', r'salary.*payable', r'working\s+hours']
            },
            'commercial_agreement': {
                'primary': ['commercial agreement', 'business agreement', 'service agreement', 'supply agreement'],
                'secondary': ['terms and conditions', 'payment terms', 'delivery', 'force majeure'],
                'patterns': [r'terms?\s+and\s+conditions', r'payment\s+terms', r'force\s+majeure']
            },
            'licensing_application': {
                'primary': ['licensing application', 'business license', 'regulatory license', 'ADGM license'],
                'secondary': ['license fee', 'regulatory requirements', 'compliance certificate'],
                'patterns': [r'licensing\s+application', r'business\s+licen[cs]e', r'regulatory\s+licen[cs]e']
            },
            'regulatory_filing': {
                'primary': ['regulatory filing', 'compliance report', 'regulatory submission', 'ADGM filing'],
                'secondary': ['regulatory authority', 'compliance requirements', 'filing deadline'],
                'patterns': [r'regulatory\s+filing', r'compliance\s+report', r'regulatory\s+submission']
            },
            'ubo_declaration': {
                'primary': ['UBO declaration', 'ultimate beneficial owner', 'beneficial ownership', 'UBO form'],
                'secondary': ['ownership structure', 'beneficial interest', 'control structure'],
                'patterns': [r'UBO\s+declaration', r'ultimate\s+beneficial\s+owner', r'beneficial\s+ownership']
            }
        }

    def _load_nlp_model(self):
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy model loaded successfully")
        except Exception as e:
            logger.warning(f"spaCy not available: {e}")
            self.nlp = None

    def _safe_str(self, value: Any) -> str:
        if isinstance(value, str):
            return value
        elif isinstance(value, (list, tuple)):
            return ' '.join(str(item) for item in value if item)
        elif value is None:
            return ""
        else:
            return str(value)

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            filename = os.path.basename(file_path)

            paragraphs = self._extract_paragraphs(doc)
            tables = self._extract_tables(doc)
            sections = self._identify_sections(paragraphs)
            document_structure = self._analyze_document_structure(paragraphs, sections)

            full_text = ' '.join([self._safe_str(p['text']) for p in paragraphs])
            section_texts = []
            for sec in sections:
                sec_title = self._safe_str(sec.get('title', ''))
                sec_text = ' '.join(self._safe_str(p['text']) for p in sec.get('paragraphs', []))
                section_texts.append(sec_title + " " + sec_text)

            stats = {
                'word_count': len(full_text.split()) if full_text else 0,
                'paragraph_count': len(paragraphs),
                'section_count': len(sections),
                'table_count': len(tables),
                'sentence_count': self._count_sentences(full_text),
                'avg_words_per_paragraph': self._calculate_avg_words_per_paragraph(paragraphs),
                'avg_sentences_per_paragraph': self._calculate_avg_sentences_per_paragraph(paragraphs),
                'character_count': len(full_text) if full_text else 0,
                'unique_words': len(set(full_text.lower().split())) if full_text else 0
            }

            try:
                file_size = os.path.getsize(file_path)
                file_modified = datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            except Exception:
                file_size = 0
                file_modified = None
            
            return {
                'metadata': {
                    'filename': filename,
                    'file_size': file_size,
                    'file_modified': file_modified,
                    'parsing_timestamp': self._get_timestamp(),
                    'parser_version': 'production_v1.0_fixed'
                },
                'structure': {
                    'paragraphs': paragraphs,
                    'tables': tables,
                    'sections': sections,
                    'section_texts': section_texts,
                    'document_structure': document_structure
                },
                'statistics': stats,
                'analysis': {
                    'text_quality': self._analyze_text_quality(paragraphs),
                    'legal_elements': self._identify_legal_elements(full_text),
                    'readability': self._calculate_readability(full_text),
                    'document_complexity': self._assess_complexity(paragraphs, sections),
                    'language_analysis': self._analyze_language(full_text)
                },
                'content': {
                    'full_text': full_text,
                    'key_phrases': self._extract_key_phrases(full_text),
                    'named_entities': self._extract_named_entities(full_text),
                    'legal_terms': self._extract_legal_terms(full_text)
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {str(e)}")
            raise Exception(f"Failed to parse document: {str(e)}")

    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        paragraphs = []
        for i, p in enumerate(doc.paragraphs):
            text = self._safe_str(p.text.strip())
            if text:
                paragraphs.append({
                    'index': i,
                    'text': text,
                    'word_count': len(text.split()) if text else 0,
                    'sentence_count': self._count_sentences(text),
                    'is_header': self._is_likely_header(text),
                    'is_list_item': self._is_list_item(text),
                    'content_type': self._classify_paragraph_type(text),
                    'legal_significance': self._assess_legal_significance(text)
                })
        return paragraphs

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        tables_data = []
        for idx, table in enumerate(doc.tables):
            try:
                rows = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = self._safe_str(cell.text.strip())
                        row_data.append(cell_text)
                    rows.append(row_data)
                
                tables_data.append({
                    'index': idx,
                    'rows': rows,
                    'row_count': len(table.rows),
                    'column_count': len(table.columns),
                    'has_headers': self._detect_table_headers(rows),
                    'table_type': self._classify_table_type(rows)
                })
            except Exception as e:
                logger.warning(f"Failed to process table {idx}: {e}")
                continue
        
        return tables_data

    def _identify_sections(self, paragraphs: List[Dict]) -> List[Dict[str, Any]]:
        sections, current = [], None
        
        for para in paragraphs:
            para_text = self._safe_str(para.get('text', ''))
            is_header = para.get('is_header', False)
            
            if is_header or self._matches_section_pattern(para_text):
                if current:
                    sections.append(current)
                current = {
                    'title': para_text,
                    'paragraphs': [],
                    'section_type': self._classify_section_type(para_text),
                    'importance': self._assess_section_importance(para_text)
                }
            
            if current:
                current['paragraphs'].append(para)
        
        if current:
            sections.append(current)
        
        return sections

    def _classify_paragraph_type(self, text: Union[str, Any]) -> str:
        text = self._safe_str(text).lower()
        
        if not text:
            return 'empty'
        
        if self._is_likely_header(text):
            return 'header'
        elif self._is_list_item(text):
            return 'list_item'
        elif len(text.split()) < 5:
            return 'short_text'
        elif any(term in text for term in ['shall', 'must', 'required', 'obligation']):
            return 'legal_obligation'
        elif any(term in text for term in ['definition', 'means', 'interpret']):
            return 'definition'
        else:
            return 'body_paragraph'

    def _assess_legal_significance(self, text: Union[str, Any]) -> str:
        text = self._safe_str(text).lower()
        legal_keywords = ['shall', 'must', 'required', 'prohibited', 'liable', 'breach', 'penalty']
        count = sum(1 for keyword in legal_keywords if keyword in text)
        
        if count >= 3:
            return 'high'
        elif count >= 1:
            return 'medium'
        else:
            return 'low'

    def _detect_table_headers(self, rows: List[List[str]]) -> bool:
        if not rows or len(rows) < 2:
            return False
        
        first_row = [self._safe_str(cell) for cell in rows[0]]
        return any(cell.isupper() or any(word.istitle() for word in cell.split()) for cell in first_row if cell)

    def _classify_table_type(self, rows: List[List[str]]) -> str:
        if not rows:
            return 'empty'
        
        headers_text = ' '.join(self._safe_str(cell).lower() for cell in rows if cell)
        
        if any(term in headers_text for term in ['name', 'director', 'member']):
            return 'personnel_list'
        elif any(term in headers_text for term in ['amount', 'fee', 'cost', 'price']):
            return 'financial_table'
        elif any(term in headers_text for term in ['date', 'deadline', 'schedule']):
            return 'schedule_table'
        else:
            return 'general_table'

    def _classify_section_type(self, title: Union[str, Any]) -> str:
        title_lower = self._safe_str(title).lower()
        
        if 'definition' in title_lower:
            return 'definitions'
        elif any(term in title_lower for term in ['article', 'clause', 'section']):
            return 'legal_provision'
        elif 'schedule' in title_lower or 'appendix' in title_lower:
            return 'appendix'
        else:
            return 'general_section'

    def _assess_section_importance(self, title: Union[str, Any]) -> str:
        title_lower = self._safe_str(title).lower()
        important_terms = ['definitions', 'objects', 'liability', 'governance', 'directors', 'shareholders']
        
        if any(term in title_lower for term in important_terms):
            return 'high'
        elif any(term in title_lower for term in ['schedule', 'appendix', 'exhibit']):
            return 'low'
        else:
            return 'medium'

    def _analyze_language(self, text: Union[str, Any]) -> Dict[str, Any]:
        text = self._safe_str(text)
        words = text.split() if text else []
        
        return {
            'average_word_length': round(sum(len(word) for word in words) / len(words), 2) if words else 0,
            'formal_language_score': self._calculate_formality_score(text),
            'passive_voice_count': len(re.findall(r'\b(?:is|was|are|were|been|being)\s+\w+ed\b', text)),
            'legal_jargon_density': self._calculate_legal_jargon_density(text)
        }

    def _calculate_formality_score(self, text: Union[str, Any]) -> float:
        text = self._safe_str(text).lower()
        formal_indicators = ['shall', 'hereby', 'whereas', 'pursuant', 'accordance', 'aforementioned']
        count = sum(1 for indicator in formal_indicators if indicator in text)
        return min(1.0, count / 10.0)

    def _calculate_legal_jargon_density(self, text: Union[str, Any]) -> float:
        text = self._safe_str(text).lower()
        legal_terms = ['jurisdiction', 'liability', 'indemnity', 'covenant', 'warranty', 'breach', 'remedy']
        word_count = len(text.split()) if text else 0
        legal_count = sum(1 for term in legal_terms if term in text)
        return round(legal_count / word_count * 100, 2) if word_count > 0 else 0

    def _extract_legal_terms(self, text: Union[str, Any]) -> List[str]:
        text = self._safe_str(text).lower()
        legal_pattern = r'\b(?:agreement|contract|party|parties|obligation|right|duty|liability|breach|remedy|jurisdiction|court|law|regulation|compliance|penalty|damages|indemnity|warranty|covenant|consideration)\b'
        return list(set(re.findall(legal_pattern, text)))

    def _analyze_document_structure(self, paragraphs, sections):
        return {
            'has_title': self._has_document_title(paragraphs),
            'has_introduction': any('introduction' in self._safe_str(s.get('title', '')).lower() for s in sections),
            'has_definitions': any('definition' in self._safe_str(s.get('title', '')).lower() for s in sections),
            'has_conclusion': bool(sections and any(w in self._safe_str(sections[-1].get('title', '')).lower() for w in ['conclusion', 'final'])),
            'section_hierarchy': {'total_levels': len(sections), 'max_level': len(sections)},
            'structural_completeness': 1.0 if len(sections) >= 3 else len(sections) / 3.0,
            'organization_score': min(1.0, sum(1 for p in paragraphs if p.get('is_header', False)) / max(1, len(paragraphs)) * 5)
        }

    def _analyze_text_quality(self, paragraphs):
        return {
            'clarity_score': 0.8,
            'completeness_indicators': ['has_sections', 'has_headers', 'structured_content'],
            'formatting_consistency': {'formatting_score': 0.85}
        }

    def _identify_legal_elements(self, text: Union[str, Any]):
        text = self._safe_str(text)
        text_lower = text.lower()
        
        return {
            'legal_patterns': {'obligations': {'count': text_lower.count('shall')}},
            'compliance_indicators': self._find_compliance_references(text),
            'contract_elements': {
                'parties': 'party' in text_lower or 'parties' in text_lower,
                'terms': 'terms' in text_lower,
                'conditions': 'conditions' in text_lower,
                'consideration': 'consideration' in text_lower
            },
            'risk_indicators': self._identify_risk_terms(text)
        }

    def _find_compliance_references(self, text: Union[str, Any]) -> List[str]:
        text = self._safe_str(text)
        compliance_patterns = [
            r'ADGM\s+[Rr]egulations?\s+\d{4}',
            r'Companies\s+Regulations\s+\d{4}',
            r'Financial\s+Services\s+and\s+Markets\s+Regulations\s+\d{4}'
        ]
        references = []
        for pattern in compliance_patterns:
            references.extend(re.findall(pattern, text))
        return references

    def _identify_risk_terms(self, text: Union[str, Any]) -> List[str]:
        text = self._safe_str(text).lower()
        risk_terms = ['penalty', 'breach', 'default', 'violation', 'non-compliance', 'liable', 'damages']
        found_terms = []
        for term in risk_terms:
            if term in text:
                found_terms.append(term)
        return found_terms

    def _calculate_readability(self, text: Union[str, Any]):
        text = self._safe_str(text)
        words = text.split() if text else []
        sentences = self._count_sentences(text)
        
        if sentences == 0:
            return {'complexity_score': 0, 'readability_level': 'unknown'}
        
        avg_sentence_length = len(words) / sentences
        complexity_score = round(avg_sentence_length * 2, 2)
        
        if complexity_score < 20:
            level = 'easy'
        elif complexity_score < 30:
            level = 'moderate'
        else:
            level = 'complex'
            
        return {'complexity_score': complexity_score, 'readability_level': level}

    def _assess_complexity(self, paragraphs, sections):
        total_words = sum(p.get('word_count', 0) for p in paragraphs)
        avg_section_length = total_words / len(sections) if sections else total_words
        
        if total_words > 3000 or avg_section_length > 500:
            complexity = 'high'
        elif total_words > 1500 or avg_section_length > 250:
            complexity = 'medium'
        else:
            complexity = 'low'
            
        return {
            'overall_complexity': complexity,
            'total_words': total_words,
            'average_section_length': round(avg_section_length, 2)
        }

    def _count_sentences(self, text: Union[str, Any]) -> int:
        text = self._safe_str(text)
        return len(re.findall(r'[.!?]+', text))
    
    def _calculate_avg_words_per_paragraph(self, pars) -> float:
        word_counts = [p.get('word_count', 0) for p in pars]
        return round(sum(word_counts) / max(1, len(pars)), 2)
    
    def _calculate_avg_sentences_per_paragraph(self, pars) -> float:
        sentence_counts = [p.get('sentence_count', 0) for p in pars]
        return round(sum(sentence_counts) / max(1, len(pars)), 2)
    
    def _is_likely_header(self, text: Union[str, Any]) -> bool:
        text = self._safe_str(text)
        return text.isupper() or bool(re.match(r'^(ARTICLE|SECTION|CLAUSE|PART|CHAPTER)\s+\d+', text, re.IGNORECASE))
    
    def _matches_section_pattern(self, text: Union[str, Any]) -> bool:
        text = self._safe_str(text)
        return any(re.match(p, text, re.IGNORECASE) for p in self.section_patterns.values())
    
    def _is_list_item(self, text: Union[str, Any]) -> bool:
        text = self._safe_str(text).strip()
        return bool(re.match(r'^\d+\.|\([a-z]\)|[•\-\*]', text))
    
    def _get_timestamp(self) -> str:
        return datetime.now().isoformat()
    
    def _has_document_title(self, paragraphs) -> bool:
        return bool(paragraphs and paragraphs[0].get('is_header', False))
    
    def _extract_key_phrases(self, text: Union[str, Any]) -> List[str]:
        text = self._safe_str(text)
        
        if self.nlp and text:
            try:
                return [chunk.text for chunk in self.nlp(text).noun_chunks][:15]
            except Exception as e:
                logger.warning(f"spaCy processing failed: {e}")
        
        return list(set(re.findall(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*', text)))[:15]
    
    def _extract_named_entities(self, text: Union[str, Any]) -> List[Dict[str, str]]:
        text = self._safe_str(text)
        
        if self.nlp and text:
            try:
                return [{'text': ent.text, 'label': ent.label_} for ent in self.nlp(text).ents][:20]
            except Exception as e:
                logger.warning(f"spaCy NER failed: {e}")

        org_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Limited|Ltd|Corporation|Corp|Company|Co)\b'
        return [{'text': match, 'label': 'ORGANIZATION'} for match in re.findall(org_pattern, text)][:20]

def parse_document(file_path: str):
    return DocumentParser().parse_document(file_path)

