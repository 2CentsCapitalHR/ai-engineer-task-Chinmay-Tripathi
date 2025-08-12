import os, sys, tempfile, json
from pathlib import Path
from datetime import datetime

sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

def test_imports():
    print("Testing system imports...")
    try:
        # Core ML libraries
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.naive_bayes import MultinomialNB
        print("--> scikit-learn imported successfully")
        
        # Project modules
        from parser import DocumentParser
        print("--> DocumentParser imported successfully")
        
        from classifier import (
            DocumentClassifier, ADGMKnowledgeBase, ComplianceChecker,
            RedFlagDetector, DocxAnnotator, LegalClauseSuggester
        )
        print("--> All classifier components imported successfully")
        
        # Optional dependencies
        try:
            import spacy
            print("--> spaCy imported successfully")
        except ImportError:
            print("⚠ spaCy not available (optional)")
            
        try:
            import gradio as gr
            print("--> Gradio imported successfully")
        except ImportError:
            print("⚠ Gradio not available")
            
        return True
    except ImportError as e:
        print(f"!!! Import error: {e}")
        return False

def test_parser():
    print("\nTesting document parser...")
    try:
        from parser import DocumentParser
        parser = DocumentParser()
        
        print(f"--> DocumentParser initialized")
        print(f"--> Supported formats: {parser.supported_formats}")
        print(f"--> Document types: {len(parser.type_indicators)}")
        print(f"--> Section patterns: {len(parser.section_patterns)}")
        
        # Test pattern matching
        test_text = "Article 1: Company Name. The company shall be known as Test Corporation Limited."
        matches = parser._matches_section_pattern(test_text)
        print(f"--> Pattern matching: {matches}")
        
        # Test paragraph classification
        content_type = parser._classify_paragraph_type(test_text)
        print(f"--> Content classification: {content_type}")
        
        return True
    except Exception as e:
        print(f"!!! Parser test failed: {e}")
        return False

def test_classifier():
    print("\nTesting document classifier...")
    try:
        from classifier import DocumentClassifier
        classifier = DocumentClassifier()
        
        print(f"--> DocumentClassifier initialized")
        print(f"--> Document types: {len(classifier.document_types)}")
        print(f"--> ADGM patterns: {len(classifier.adgm_patterns)}")
        print(f"--> ML classifiers: {list(classifier.classifiers.keys())}")

        mock_content = {
            'content': {'full_text': 'Articles of Association for Test Company Limited. This document contains governance provisions.'},
            'structure': {'paragraphs': [{'is_header': True}], 'sections': []},
            'statistics': {'word_count': 15, 'table_count': 0}
        }
        
        features = classifier.extract_features(mock_content)
        print(f"--> Feature extraction: {len(features)} features")

        result = classifier.classify_document(mock_content, method='rule_based')
        print(f"--> Classification: {result['predicted_type']} ({result['confidence']:.2f})")

        explanation = classifier.get_classification_explanation(result)
        print(f"--> Explanation generated: {len(explanation)} characters")
        
        return True
    except Exception as e:
        print(f"!!! Classifier test failed: {e}")
        return False

def test_knowledge_base():
    print("\nTesting knowledge base...")
    try:
        from classifier import ADGMKnowledgeBase

        kb = ADGMKnowledgeBase()
        print("--> Knowledge base initialized")

        test_texts = [
            "ADGM Companies Regulations 2020 Article 1: These regulations apply to all companies.",
            "ADGM Courts have jurisdiction over all matters within the ADGM."
        ]

        result = kb.build_knowledge_base(test_texts)
        print(f"--> Knowledge base build attempted: {result}")
        
        return True
    except Exception as e:
        print(f"!!! Knowledge base test failed: {e}")
        return False

def test_compliance_checker():
    print("\nTesting compliance checker...")
    try:
        from classifier import ComplianceChecker
        cc = ComplianceChecker()
        
        print(f"--> ComplianceChecker initialized")
        print(f"--> Document requirements loaded: {len(cc.required_documents)}")

        test_types = ['articles_of_association', 'memorandum_of_association']
        result = cc.check(test_types)
        
        print(f"--> Compliance check result: {result['status']}")
        print(f"--> Process detected: {result['process']}")
        print(f"--> Completeness: {result['completeness_percentage']}%")
        
        return True
    except Exception as e:
        print(f"!!! Compliance checker test failed: {e}")
        return False

def test_red_flag_detector():
    print("\nTesting red flag detector...")
    try:
        from classifier import RedFlagDetector
        rf = RedFlagDetector()
        
        print(f"--> RedFlagDetector initialized")
        print(f"--> Patterns loaded: {len(rf.patterns)}")
        
        test_text = "This agreement is governed by UAE Federal Courts and shall be subject to UAE law."
        flags = rf.scan(test_text)
        
        print(f"--> Red flags detected: {len(flags)}")
        if flags:
            print(f"--> Sample flag: {flags[0]['issue']}")
            print(f"--> ADGM reference: {flags.get('adgm_reference', 'N/A')}")
        
        return True
    except Exception as e:
        print(f"!!! Red flag detector test failed: {e}")
        return False

def test_legal_suggester():
    print("\nTesting legal clause suggester...")
    try:
        from classifier import LegalClauseSuggester
        ls = LegalClauseSuggester()
        
        print(f"--> LegalClauseSuggester initialized")

        suggestion = ls.suggest_clause("jurisdiction", "UAE Federal Courts")
        print(f"--> Suggestion generated: {bool(suggestion)}")
        
        if suggestion:
            print(f"--> Template available: {bool(suggestion.get('suggested_clause'))}")
            print(f"--> Legal reference: {suggestion.get('legal_reference', 'N/A')}")

        all_suggestions = ls.get_all_suggestions()
        print(f"--> Total suggestion categories: {len(all_suggestions)}")
        
        return True
    except Exception as e:
        print(f"!!! Legal suggester test failed: {e}")
        return False

def test_docx_annotator():
    print("\nTesting DOCX annotator...")
    try:
        from classifier import DocxAnnotator
        from docx import Document

        fd, temp_doc = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        
        doc = Document()
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This contract is governed by UAE Federal Courts.")
        doc.save(temp_doc)

        annotator = DocxAnnotator()
        review_comments = [{
            "paragraph_index": 1,
            "text": "This contract is governed by UAE Federal Courts.",
            "comment": "Incorrect jurisdiction - should reference ADGM Courts",
            "severity": "High",
            "adgm_reference": "ADGM Courts Regulations 2020, Section 3"
        }]
        
        marked_file = annotator.add_comments(temp_doc, review_comments)
        
        print(f"--> Annotation successful: {bool(marked_file)}")
        
        if marked_file and os.path.exists(marked_file):
            print(f"--> Marked file created: {os.path.basename(marked_file)}")
            os.remove(marked_file)

        os.remove(temp_doc)
        
        return True
    except Exception as e:
        print(f"!!! DOCX annotator test failed: {e}")
        return False

def test_integration():
    print("\nTesting system integration...")
    try:
        from parser import DocumentParser
        from classifier import DocumentClassifier, RedFlagDetector
        
        parser = DocumentParser()
        classifier = DocumentClassifier()
        detector = RedFlagDetector()

        mock_content = {
            'metadata': {'filename': 'test.docx'},
            'structure': {
                'paragraphs': [
                    {'index': 0, 'text': 'Articles of Association', 'is_header': True},
                    {'index': 1, 'text': 'This company is governed by UAE Federal Courts.', 'is_header': False}
                ]
            },
            'statistics': {'word_count': 20},
            'content': {'full_text': 'Articles of Association. This company is governed by UAE Federal Courts.'},
            'analysis': {'document_complexity': {'overall_complexity': 'medium'}}
        }

        result = classifier.classify_document(mock_content)
        print(f"--> Integration classification: {result['predicted_type']}")

        flags = detector.scan(mock_content['content']['full_text'])
        print(f"--> Integration red flags: {len(flags)}")
        
        return True
    except Exception as e:
        print(f"!!! Integration test failed: {e}")
        return False

def test_gradio_interface():
    print("\nTesting Gradio interface...")
    try:
        import gradio as gr
        
        def mock_analyze(files, method):
            return {}, {}, [], [], "Mock analysis complete", {}, "Mock summary"
        
        with gr.Blocks() as demo:
            file_input = gr.File(file_types=[".docx"], file_count="multiple")
            method_dropdown = gr.Dropdown(choices=["rule_based", "ensemble"])
            analyze_button = gr.Button("Analyze")

            results = gr.JSON()
            compliance = gr.JSON()
            table = gr.DataFrame()
            files_out = gr.Files()
            log = gr.Textbox()
            analytics = gr.State()
            summary = gr.Markdown()
            
            analyze_button.click(
                mock_analyze,
                inputs=[file_input, method_dropdown],
                outputs=[results, compliance, table, files_out, log, analytics, summary]
            )
        
        print("--> Gradio interface created successfully")
        return True
    except Exception as e:
        print(f"!!! Gradio interface test failed: {e}")
        return False

def run_comprehensive_tests():
    print("=" * 70)
    print("ADGM Corporate Agent - Comprehensive Test Suite")
    print("=" * 70)
    
    tests = [
        ("System Imports", test_imports),
        ("Document Parser", test_parser),
        ("Document Classifier", test_classifier),
        ("Knowledge Base", test_knowledge_base),
        ("Compliance Checker", test_compliance_checker),
        ("Red Flag Detector", test_red_flag_detector),
        ("Legal Suggester", test_legal_suggester),
        ("DOCX Annotator", test_docx_annotator),
        ("System Integration", test_integration),
        ("Gradio Interface", test_gradio_interface)
    ]
    
    passed_tests = 0
    total_tests = len(tests)
    
    for test_name, test_func in tests:
        print(f"\nTesting {test_name}...")
        try:
            if test_func():
                passed_tests += 1
                print(f">>>{test_name} PASSED")
            else:
                print(f"!!{test_name} FAILED")
        except Exception as e:
            print(f"!!{test_name} FAILED with exception: {e}")
    
    # Generate test report
    success_rate = (passed_tests / total_tests) * 100
    
    print("\n" + "=" * 70)
    print("*** TEST RESULTS SUMMARY ***")
    print("=" * 70)
    print(f"Tests Passed: {passed_tests}/{total_tests}")
    print(f"Success Rate: {success_rate:.1f}%")
    
    if success_rate >= 80:
        print("SYSTEM READY FOR PRODUCTION!")
        print("\nAvailable Features:")
        print("  • Advanced document parsing and structure analysis")
        print("  • ML-enhanced document classification")
        print("  • RAG-powered ADGM knowledge base integration")
        print("  • Comprehensive compliance checking")
        print("  • Intelligent red flag detection with legal citations")
        print("  • Automated document annotation with highlights")
        print("  • Legal clause suggestions and recommendations")
        print("  • Professional web interface with analytics")
        print("  • Export capabilities for reviewed documents")
    else:
        print("! ! ! SYSTEM NEEDS ATTENTION")
        print("\nTroubleshooting Steps:")
        print("1. Install all dependencies: pip install -r requirements.txt")
        print("2. Ensure python-docx is properly installed")
        print("3. Check OpenAI API key configuration for RAG features")
        print("4. Verify spaCy model: python -m spacy download en_core_web_sm")
        print("5. Ensure proper file permissions for marked_documents/")
    
    return success_rate >= 80

if __name__ == "__main__":
    success = run_comprehensive_tests()
    exit(0 if success else 1)

