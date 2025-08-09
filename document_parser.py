from docx import Document
import tempfile
import os
from typing import Dict, List, Any
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentParser:
    
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        
        try:
            doc = Document(file_path)
            
            content = {
                'filename': os.path.basename(file_path),
                'paragraphs': self._extract_paragraphs(doc),
                'tables': self._extract_tables(doc),
                'sections': self._identify_sections(doc),
                'word_count': self._count_words(doc),
                'paragraph_count': len(doc.paragraphs)
            }
            
            logger.info(f"---Successfully parsed document: {content['filename']}")
            return content
            
        except Exception as e:
            logger.error(f"!!!Error parsing document {file_path}: {str(e)}")
            raise Exception(f"!!!Failed to parse document: {str(e)}")
    
    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Skipping empty paragraphs
                paragraphs.append({
                    'index': i,
                    'text': paragraph.text.strip(),
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'word_count': len(paragraph.text.split())
                })
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[List[List[str]]]:
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def _identify_sections(self, doc: Document) -> List[Dict[str, Any]]:
        sections = []
        current_section = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else 'Normal'
            
            # Check if this might be a section header
            if self._is_section_header(text, style):
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'title': text,
                    'start_paragraph': i,
                    'style': style,
                    'content': []
                }
            elif current_section and text:
                current_section['content'].append({
                    'paragraph_index': i,
                    'text': text
                })
        
        # Important last section
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _is_section_header(self, text: str, style: str) -> bool:
        if 'Heading' in style or 'Title' in style:
            return True
        
        section_indicators = [
            'article', 'clause', 'section', 'part', 'chapter',
            'memorandum', 'whereas', 'resolved', 'definitions'
        ]
        
        text_lower = text.lower()
        if any(indicator in text_lower for indicator in section_indicators):
            return True

        if len(text.split()) <= 10 and text.isupper():
            return True
        
        return False
    
    def _count_words(self, doc: Document) -> int:
        total_words = 0
        for paragraph in doc.paragraphs:
            total_words += len(paragraph.text.split())
        return total_words
    
    def get_document_type_hints(self, content: Dict[str, Any]) -> List[str]:
        
        hints = []
        full_text = ' '.join([p['text'] for p in content['paragraphs']]).lower()

        type_indicators = {
            'articles_of_association': [
                'articles of association', 'governance', 'directors', 'shares',
                'meetings', 'board', 'quorum'
            ],
            'memorandum_of_association': [
                'memorandum of association', 'objects', 'liability limited',
                'company limited', 'purposes'
            ],
            'incorporation_application': [
                'incorporation', 'application', 'registration', 'company formation',
                'registration authority'
            ],
            'board_resolution': [
                'resolution', 'resolved', 'board meeting', 'directors resolve',
                'unanimous consent'
            ],
            'register_members': [
                'register of members', 'shareholders', 'member details',
                'share certificate', 'membership'
            ]
        }
        
        for doc_type, keywords in type_indicators.items():
            matches = sum(1 for keyword in keywords if keyword in full_text)
            if matches >= 2:  # At least 2 keyword matches
                hints.append(doc_type)
        
        return hints if hints else ['unknown_document_type']

def parse_uploaded_file(file_path: str) -> Dict[str, Any]:

    parser = DocumentParser()
    return parser.parse_document(file_path)
