import re, os, logging
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import SVC
from sklearn.ensemble import RandomForestClassifier

from langchain_community.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter

from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

class DocumentClassifier:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(
            max_features=5000, stop_words='english', ngram_range=(1, 3),
            min_df=2, max_df=0.95
        )
        
        self.classifiers = {
            'naive_bayes': MultinomialNB(alpha=0.1),
            'svm': SVC(kernel='linear', probability=True),
            'random_forest': RandomForestClassifier(n_estimators=100, random_state=42)
        }
        
        self.is_trained = False
        self.document_types = [
            'articles_of_association', 'memorandum_of_association', 'incorporation_application',
            'board_resolution', 'register_members_directors', 'employment_contract',
            'commercial_agreement', 'licensing_application', 'regulatory_filing',
            'compliance_policy', 'ubo_declaration', 'shareholder_resolution',
            'power_of_attorney', 'audit_report', 'financial_statements'
        ]

        self.adgm_patterns = {
            'articles_of_association': {
                'primary_keywords': [
                    'articles of association', 'company governance', 'board of directors',
                    'shareholders', 'general meeting', 'quorum', 'voting rights',
                    'share capital', 'dividend distribution', 'board composition'
                ],
                'secondary_keywords': [
                    'annual general meeting', 'extraordinary resolution', 'ordinary resolution',
                    'director appointment', 'company secretary', 'audit committee',
                    'remuneration committee', 'corporate governance', 'shareholder rights'
                ],
                'legal_patterns': [
                    r'article\s+\d+', r'board\s+of\s+directors', r'general\s+meeting',
                    r'share\s+capital', r'voting\s+rights', r'dividend\s+policy'
                ],
                'section_indicators': ['governance', 'meetings', 'directors', 'shares', 'capital']
            },
            'memorandum_of_association': {
                'primary_keywords': [
                    'memorandum of association', 'company objects', 'liability limited',
                    'registered office', 'authorized capital', 'company formation'
                ],
                'secondary_keywords': [
                    'company purposes', 'business activities', 'share classes',
                    'incorporation', 'registered address', 'capital structure'
                ],
                'legal_patterns': [
                    r'objects?\s+of\s+(?:the\s+)?company', r'liability.*limited',
                    r'registered\s+office', r'authorized\s+capital', r'share\s+classes'
                ],
                'section_indicators': ['objects', 'liability', 'capital', 'office', 'purposes']
            },
            'ubo_declaration': {
                'primary_keywords': [
                    'UBO declaration', 'ultimate beneficial owner', 'beneficial ownership',
                    'ownership structure', 'control structure', 'beneficial interest'
                ],
                'secondary_keywords': [
                    'ownership percentage', 'voting control', 'economic interest',
                    'trust arrangement', 'nominee arrangement'
                ],
                'legal_patterns': [
                    r'UBO\s+declaration', r'ultimate\s+beneficial\s+owner',
                    r'beneficial\s+ownership', r'\d+%\s+(?:ownership|interest)'
                ],
                'section_indicators': ['ownership', 'control', 'beneficial', 'ultimate']
            }
        }

    def _safe_get(self, obj: Any, key: str, default: Any = None) -> Any:
        if isinstance(obj, dict):
            return obj.get(key, default)
        elif hasattr(obj, key):
            return getattr(obj, key, default)
        else:
            return default

    def extract_features(self, document_content):
        content = self._safe_get(document_content, 'content', {})
        full_text = self._safe_get(content, 'full_text', '') if isinstance(content, dict) else ''
        
        structure = self._safe_get(document_content, 'structure', {})
        paragraphs = self._safe_get(structure, 'paragraphs', []) if isinstance(structure, dict) else []
        sections = self._safe_get(structure, 'sections', []) if isinstance(structure, dict) else []
        
        statistics = self._safe_get(document_content, 'statistics', {})
        
        return {
            'full_text': str(full_text) if full_text else '',
            'word_count': self._safe_get(statistics, 'word_count', 0) if isinstance(statistics, dict) else 0,
            'paragraph_count': len(paragraphs) if isinstance(paragraphs, list) else 0,
            'section_count': len(sections) if isinstance(sections, list) else 0,
            'has_title': self._has_document_title(paragraphs),
            'has_sections': len(sections) > 0 if isinstance(sections, list) else False,
            'has_numbered_sections': self._has_numbered_sections(sections),
            'has_tables': self._safe_get(statistics, 'table_count', 0) > 0 if isinstance(statistics, dict) else False,
            'legal_term_density': self._calculate_legal_term_density(full_text),
            'formality_score': self._calculate_formality_score(full_text)
        }

    def classify_document(self, document_content, method: str = 'ensemble'):
        features = self.extract_features(document_content)
        
        if method == 'rule_based':
            return self._classify_rule_based(features)
        elif method == 'ml' and self.is_trained:
            return self._classify_ml(features)
        elif method == 'ensemble':
            return self._classify_ensemble(features)
        else:
            return self._classify_rule_based(features)

    def _classify_rule_based(self, features):
        text = str(features.get('full_text', '')).lower()
        scores = {}
        
        for doc_type, patterns in self.adgm_patterns.items():
            score = 0.0
            matches = {'primary': [], 'secondary': [], 'patterns': [], 'sections': []}

            primary_keywords = patterns.get('primary_keywords', [])
            for keyword in primary_keywords:
                if str(keyword).lower() in text:
                    score += 3.0
                    matches['primary'].append(keyword)

            secondary_keywords = patterns.get('secondary_keywords', [])
            for keyword in secondary_keywords:
                if str(keyword).lower() in text:
                    score += 1.5
                    matches['secondary'].append(keyword)

            legal_patterns = patterns.get('legal_patterns', [])
            for pattern in legal_patterns:
                if re.search(str(pattern), text, re.IGNORECASE):
                    score += 2.5
                    matches['patterns'].append(pattern)

            section_indicators = patterns.get('section_indicators', [])
            for indicator in section_indicators:
                if str(indicator).lower() in text:
                    score += 1.0
                    matches['sections'].append(indicator)

            text_length_factor = len(text.split()) / 1000.0 if text.split() else 1.0
            normalized_score = score / max(text_length_factor, 0.1)
            
            scores[doc_type] = {
                'raw_score': score,
                'normalized_score': normalized_score,
                'matches': matches
            }
        
        if not scores:
            return {
                'predicted_type': 'unknown',
                'confidence': 0.0,
                'method_used': 'rule_based',
                'all_scores': {}
            }
        
        best_type = max(scores.keys(), key=lambda x: scores[x]['normalized_score'])
        best_score = scores[best_type]['normalized_score']
        sorted_scores = sorted(scores.values(), key=lambda x: x['normalized_score'], reverse=True)

        confidence = min(1.0, best_score / 15.0)
        if len(sorted_scores) > 1:
            score_separation = sorted_scores[0]['normalized_score'] - sorted_scores[1]['normalized_score']
            confidence += min(0.4, score_separation / 10.0)
        confidence = min(1.0, confidence)
        
        return {
            'predicted_type': best_type,
            'confidence': round(confidence, 3),
            'method_used': 'rule_based',
            'all_scores': {k: round(v['normalized_score'], 3) for k, v in scores.items()},
            'best_matches': scores[best_type]['matches'],
            'reasoning': f"Matched {len(scores[best_type]['matches']['primary'])} primary keywords, "
                        f"{len(scores[best_type]['matches']['patterns'])} patterns"
        }

    def _classify_ensemble(self, features):
        rule_result = self._classify_rule_based(features)
        rule_result['method_used'] = 'ensemble'
        
        confidence = rule_result['confidence']
        matches = rule_result.get('best_matches', {})
        evidence_types = sum(1 for match_type in matches.values() if match_type)

        if evidence_types >= 3:
            confidence = min(1.0, confidence * 1.3)
        elif evidence_types >= 2:
            confidence = min(1.0, confidence * 1.1)
        
        rule_result['confidence'] = round(confidence, 3)
        rule_result['evidence_types'] = evidence_types
        
        return rule_result

    def get_classification_explanation(self, classification_result):
        predicted_type = classification_result.get('predicted_type', 'unknown')
        confidence = classification_result.get('confidence', 0.0)
        method = classification_result.get('method_used', 'unknown')
        
        explanation = f"Document classified as '{predicted_type.replace('_', ' ').title()}' "
        explanation += f"with {confidence:.1%} confidence using {method} method.\n\n"
        
        if 'best_matches' in classification_result:
            matches = classification_result['best_matches']
            if matches.get('primary'):
                explanation += f"Primary indicators: {', '.join(matches['primary'][:3])}\n"
            if matches.get('patterns'):
                explanation += f"Legal patterns: {len(matches['patterns'])} patterns matched\n"
            if matches.get('secondary'):
                explanation += f"Secondary indicators: {len(matches['secondary'])} found\n"
        
        if 'reasoning' in classification_result:
            explanation += f"\n{classification_result['reasoning']}"
        
        return explanation

    def _calculate_legal_term_density(self, text: Union[str, Any]) -> float:
        text_str = str(text) if text else ''
        legal_terms = ['shall', 'hereby', 'whereas', 'pursuant', 'liability', 'obligation']
        word_count = len(text_str.split())
        legal_count = sum(1 for term in legal_terms if term in text_str.lower())
        return round(legal_count / word_count * 100, 2) if word_count > 0 else 0

    def _calculate_formality_score(self, text: Union[str, Any]) -> float:
        text_str = str(text) if text else ''
        formal_indicators = ['shall', 'hereby', 'whereas', 'pursuant', 'accordance', 'aforementioned']
        count = sum(1 for indicator in formal_indicators if indicator in text_str.lower())
        return min(1.0, count / 15.0)

    def _has_document_title(self, paragraphs):
        if not isinstance(paragraphs, list) or not paragraphs:
            return False
        
        first_para = paragraphs[0]
        if isinstance(first_para, dict):
            return first_para.get('is_header', False) or first_para.get('content_type', '') == 'title_header'
        
        return False

    def _has_numbered_sections(self, sections):
        if not isinstance(sections, list):
            return False
        
        for section in sections:
            if isinstance(section, dict):
                title = section.get('title', '')
                if re.match(r'^\d+\.', str(title)):
                    return True
        
        return False

    def _classify_ml(self, features):
        return self._classify_rule_based(features)  # Fallback until ML is trained

# RAG
class ADGMKnowledgeBase:
    def __init__(self, openai_api_key=None):
        self.embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key) if openai_api_key else None
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        self.vectorstore = None
        self.knowledge_index = {}

    def build_knowledge_base(self, list_of_texts):
        if not self.embeddings:
            logger.warning("OpenAI API key not provided for embeddings")
            return False
        
        try:
            chunks = []
            for i, txt in enumerate(list_of_texts):
                doc_chunks = self.text_splitter.split_text(str(txt))
                for chunk in doc_chunks:
                    chunks.append(chunk)
                    self.knowledge_index[len(chunks)-1] = f"document_{i}"
            
            self.vectorstore = FAISS.from_texts(chunks, self.embeddings)
            logger.info(f"Knowledge base built with {len(chunks)} chunks from {len(list_of_texts)} documents")
            return True
        except Exception as e:
            logger.error(f"Failed to build knowledge base: {e}")
            return False

    def retrieve(self, query, k=3):
        if not self.vectorstore:
            return []
        
        try:
            docs = self.vectorstore.similarity_search(str(query), k=k)
            return [{'content': doc.page_content, 'relevance': 'high'} for doc in docs]
        except Exception as e:
            logger.error(f"Failed to retrieve from knowledge base: {e}")
            return []

class ComplianceChecker:
    def __init__(self):
        self.required_documents = {
            "company_incorporation": [
                "articles_of_association",
                "memorandum_of_association", 
                "incorporation_application",
                "register_members_directors",
                "board_resolution"
            ],
            "licensing_application": [
                "licensing_application",
                "articles_of_association",
                "memorandum_of_association",
                "regulatory_filing"
            ],
            "compliance_filing": [
                "compliance_policy",
                "regulatory_filing",
                "audit_report"
            ]
        }
        
        self.document_requirements = {
            "articles_of_association": [
                "company governance structure",
                "board composition and duties", 
                "shareholder rights and meetings",
                "share capital provisions"
            ],
            "memorandum_of_association": [
                "company objects and purposes",
                "liability limitation clauses",
                "registered office address",
                "authorized share capital"
            ]
        }

    def check(self, document_types_list):
        if not isinstance(document_types_list, list):
            document_types_list = []

        process = self._determine_process(document_types_list)
        required_docs = self.required_documents.get(process, [])
        
        present_docs = set(str(doc) for doc in document_types_list)
        required_docs_set = set(required_docs)
        missing_docs = required_docs_set - present_docs
        extra_docs = present_docs - required_docs_set
        
        completeness_percentage = len(present_docs.intersection(required_docs_set)) / len(required_docs_set) * 100 if required_docs_set else 100
        
        return {
            "process": process,
            "status": "Complete" if not missing_docs else "Incomplete",
            "completeness_percentage": round(completeness_percentage, 1),
            "total_required": len(required_docs),
            "documents_provided": len(document_types_list),
            "documents_matched": len(present_docs.intersection(required_docs_set)),
            "missing_documents": list(missing_docs),
            "extra_documents": list(extra_docs),
            "recommendations": self._generate_recommendations(missing_docs, process)
        }

    def _determine_process(self, document_types):
        if not isinstance(document_types, list):
            return "general_submission"
        
        document_types_str = [str(doc) for doc in document_types]
        
        incorporation_indicators = ['articles_of_association', 'memorandum_of_association', 'incorporation_application']
        licensing_indicators = ['licensing_application', 'regulatory_filing']
        compliance_indicators = ['compliance_policy', 'regulatory_filing', 'audit_report']
        
        incorporation_count = sum(1 for doc in document_types_str if doc in incorporation_indicators)
        licensing_count = sum(1 for doc in document_types_str if doc in licensing_indicators)
        compliance_count = sum(1 for doc in document_types_str if doc in compliance_indicators)
        
        if incorporation_count >= 2:
            return "company_incorporation"
        elif licensing_count >= 1:
            return "licensing_application"
        elif compliance_count >= 1:
            return "compliance_filing"
        else:
            return "general_submission"

    def _generate_recommendations(self, missing_docs, process):
        recommendations = []
        for doc in missing_docs:
            doc_name = str(doc).replace('_', ' ').title()
            recommendations.append(f"Please provide {doc_name} to complete your {process.replace('_', ' ')} submission")
        return recommendations

class RedFlagDetector:
    def __init__(self):
        self.patterns = [
            {
                "pattern": r"UAE Federal Courts?",
                "issue": "Incorrect jurisdiction reference",
                "severity": "High",
                "suggestion": "Replace with ADGM Courts",
                "adgm_reference": "ADGM Courts Regulations 2020, Section 3",
                "category": "jurisdiction"
            },
            {
                "pattern": r"without\s+ADGM\s+approval",
                "issue": "Missing ADGM approval requirement",
                "severity": "Medium", 
                "suggestion": "Add explicit ADGM approval clause",
                "adgm_reference": "ADGM Companies Regulations 2020, Article 15",
                "category": "regulatory_compliance"
            },
            {
                "pattern": r"UAE\s+(?:law|laws|legal system)",
                "issue": "Incorrect legal system reference",
                "severity": "High",
                "suggestion": "Reference ADGM law and regulations",
                "adgm_reference": "ADGM Legal Framework",
                "category": "jurisdiction"
            },
            {
                "pattern": r"(?:subject to|governed by)\s+UAE\s+(?:law|legislation)",
                "issue": "Incorrect governing law clause",
                "severity": "High", 
                "suggestion": "Specify ADGM law as governing law",
                "adgm_reference": "ADGM Companies Regulations 2020, Article 2",
                "category": "governing_law"
            },
            {
                "pattern": r"(?:registered|incorporated)\s+in\s+(?:UAE|United Arab Emirates)(?!\s+(?:ADGM|Abu Dhabi Global Market))",
                "issue": "Ambiguous jurisdiction for registration",
                "severity": "Medium",
                "suggestion": "Clarify registration in ADGM specifically", 
                "adgm_reference": "ADGM Registration Regulations 2020",
                "category": "registration"
            },
            {
                "pattern": r"shall\s+be\s+void",
                "issue": "Potentially harsh penalty clause",
                "severity": "Medium",
                "suggestion": "Consider graduated penalties or remedies",
                "adgm_reference": "ADGM Contract Law",
                "category": "contract_terms"
            }
        ]

    def scan(self, text):
        text_str = str(text) if text else ""
        flags = []
        
        for pattern_info in self.patterns:
            if not isinstance(pattern_info, dict):
                continue
                
            pattern = pattern_info.get("pattern", "")
            if not pattern:
                continue
                
            try:
                matches = list(re.finditer(pattern, text_str, re.IGNORECASE))
                for match in matches:
                    flag = {
                        "pattern": pattern,
                        "issue": pattern_info.get("issue", "Unknown issue"),
                        "severity": pattern_info.get("severity", "Medium"),
                        "suggestion": pattern_info.get("suggestion", "Review required"),
                        "adgm_reference": pattern_info.get("adgm_reference", ""),
                        "category": pattern_info.get("category", "general"),
                        "matched_text": match.group(),
                        "position": match.span()
                    }
                    flags.append(flag)
            except Exception as e:
                logger.warning(f"Error processing pattern {pattern}: {e}")
                continue
        
        return flags

class DocxAnnotator:
    def __init__(self):
        self.annotation_styles = {
            'high_severity': {'highlight': WD_COLOR_INDEX.RED, 'font_color': RGBColor(139, 0, 0)},
            'medium_severity': {'highlight': WD_COLOR_INDEX.YELLOW, 'font_color': RGBColor(184, 134, 11)},
            'low_severity': {'highlight': WD_COLOR_INDEX.GREEN, 'font_color': RGBColor(0, 100, 0)}
        }

    def add_comments(self, original_file_path, paragraphs_and_comments):
        try:
            doc = DocxDocument(original_file_path)
            
            for pc in paragraphs_and_comments:
                if not isinstance(pc, dict):
                    continue
                    
                para_idx = pc.get('paragraph_index', -1)
                comment = pc.get('comment', '')
                severity = str(pc.get('severity', 'medium')).lower()
                
                if para_idx >= 0 and para_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_idx]
                    
                    # Different highlights for different severity levels
                    style = self.annotation_styles.get(f'{severity}_severity', 
                                                     self.annotation_styles['medium_severity'])
                    
                    for run in paragraph.runs:
                        if run.text.strip():
                            run.font.highlight_color = style['highlight']
                    
                    legal_ref = pc.get('adgm_reference', '')
                    full_comment = f" [ADGM Review: {comment}"
                    if legal_ref:
                        full_comment += f" | Legal Reference: {legal_ref}"
                    full_comment += "]"
                    
                    comment_run = paragraph.add_run(full_comment)
                    comment_run.italic = True
                    comment_run.font.color.rgb = style['font_color']

            marked_dir = 'marked_documents'
            if not os.path.exists(marked_dir):
                os.makedirs(marked_dir)

            base_name = os.path.splitext(os.path.basename(original_file_path))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            marked_file = os.path.join(marked_dir, f"{base_name}_REVIEWED_{timestamp}.docx")
            
            doc.save(marked_file)
            logger.info(f"Annotated document saved: {marked_file}")
            return marked_file
            
        except Exception as e:
            logger.error(f"Failed to annotate document: {e}")
            return ""

class LegalClauseSuggester:
    def __init__(self):
        self.clause_templates = {
            "jurisdiction": {
                "template": "This Agreement shall be governed by and construed in accordance with the laws of the Abu Dhabi Global Market (ADGM), and the parties hereby submit to the exclusive jurisdiction of the ADGM Courts.",
                "adgm_reference": "ADGM Courts Regulations 2020"
            },
            "governing_law": {
                "template": "This document and any disputes arising out of or in connection with it shall be governed by ADGM law.",
                "adgm_reference": "ADGM Legal Framework"
            },
            "regulatory_compliance": {
                "template": "The Company shall comply with all applicable ADGM regulations and shall obtain all necessary approvals from the relevant ADGM authorities before undertaking any regulated activities.",
                "adgm_reference": "ADGM Companies Regulations 2020"
            }
        }

    def suggest_clause(self, category: str, context: str = "") -> Dict[str, str]:
        if category in self.clause_templates:
            template = self.clause_templates[category]
            return {
                "suggested_clause": template["template"],
                "legal_reference": template["adgm_reference"],
                "category": category,
                "context_note": f"Suggested replacement for: {context}" if context else ""
            }
        return {}

    def get_all_suggestions(self) -> Dict[str, Dict[str, str]]:
        return self.clause_templates

def classify_document(document_content, method: str = 'ensemble'):
    classifier = DocumentClassifier()
    return classifier.classify_document(document_content, method)

